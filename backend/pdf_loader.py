import fitz  # PyMuPDF
import re
import io
from docx import Document
from PIL import Image
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

OCR_MIN_CHARS = 20
OCR_DPI = 200


def _ocr_page(page):
    """Renders a PDF page as an image and runs OCR on it (for scanned pages,
    certificates, or any page where text is embedded as an image)."""
    zoom = OCR_DPI / 72
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    try:
        return pytesseract.image_to_string(img)
    except Exception:
        return ""


def extract_text_from_pdf(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)

    for page in doc:
        page_text = page.get_text()
        if len(page_text.strip()) < OCR_MIN_CHARS:
            page_text = _ocr_page(page)
        text += page_text

    doc.close()
    return text


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join(para.text for para in doc.paragraphs)
    return text


def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_file(file_path):
    """
    Detects the file type from its extension and routes to the right extractor.
    """
    if file_path.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif file_path.lower().endswith(".txt"):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")


def extract_pages_from_pdf(pdf_path):
    """Returns list of (page_number, text) — 1-indexed.
    Pages with little/no extractable text (scanned pages, certificates,
    image-only pages) automatically fall back to OCR."""
    pages = []
    doc = fitz.open(pdf_path)
    for i, page in enumerate(doc):
        page_text = page.get_text()
        used_ocr = False
        if len(page_text.strip()) < OCR_MIN_CHARS:
            ocr_text = _ocr_page(page)
            if len(ocr_text.strip()) > len(page_text.strip()):
                page_text = ocr_text
                used_ocr = True
        pages.append((i + 1, page_text))
        if used_ocr:
            print(f"[pdf_loader] Page {i + 1}: used OCR fallback ({len(page_text.strip())} chars).")
    doc.close()
    return pages


def extract_pages_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join(para.text for para in doc.paragraphs)
    return [(None, text)]  # docx mein reliable page number nahi hota


def extract_pages_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        return [(None, f.read())]


def extract_pages_from_file(file_path):
    """
    Har file type ko (page_number, text) tuples ki list mein return karta hai.
    """
    if file_path.lower().endswith(".pdf"):
        return extract_pages_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return extract_pages_from_docx(file_path)
    elif file_path.lower().endswith(".txt"):
        return extract_pages_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")


def clean_text(text):
    text = re.sub(r"\s+", " ", text)
    return text.strip()